import os
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

OUT = os.path.dirname(os.path.abspath(__file__))
PDF_DIR = os.path.join(OUT, "pdf")
DOCX_DIR = os.path.join(OUT, "docx")
HTML_DIR = os.path.join(OUT, "html")
MD_DIR = os.path.join(OUT, "markdown")

# ---------------------------------------------------------------------------
# Fictional builders / projects (synthetic but internally consistent data,
# so retrieval and citations behave realistically)
# ---------------------------------------------------------------------------
BUILDERS = {
    "skyline": {
        "name": "Skyline Horizon Developers",
        "founded": 1998,
        "hq": "Bengaluru, Karnataka",
        "projects": [
            {
                "id": "sht", "name": "Skyline Horizon Towers", "type": "Residential (2/3 BHK apartments)",
                "location": "Whitefield, Bengaluru", "units": 480, "towers": 4,
                "price_range": "INR 78 lakh - 1.65 crore", "possession": "December 2027",
                "rera": "PRM/KA/RERA/1251/446/PR/210324/006712",
            },
            {
                "id": "hbp", "name": "Horizon Business Park", "type": "Commercial (office spaces)",
                "location": "Outer Ring Road, Bengaluru", "units": 120, "towers": 2,
                "price_range": "INR 1.2 crore - 4.8 crore", "possession": "March 2028",
                "rera": "PRM/KA/RERA/1251/446/PR/210324/006713",
            },
        ],
    },
    "meridian": {
        "name": "Meridian Greens Realty",
        "founded": 2005,
        "hq": "Pune, Maharashtra",
        "projects": [
            {
                "id": "mgr", "name": "Meridian Greens Residency", "type": "Residential (1/2/3 BHK apartments)",
                "location": "Hinjewadi, Pune", "units": 620, "towers": 6,
                "price_range": "INR 52 lakh - 1.35 crore", "possession": "June 2027",
                "rera": "P52100034521",
            },
            {
                "id": "mlv", "name": "Meridian Lakeview Villas", "type": "Residential (4 BHK villas)",
                "location": "Baner, Pune", "units": 84, "towers": 0,
                "price_range": "INR 2.4 crore - 3.9 crore", "possession": "September 2026",
                "rera": "P52100034899",
            },
        ],
    },
    "urbannest": {
        "name": "Urban Nest Infrastructures",
        "founded": 2011,
        "hq": "Hyderabad, Telangana",
        "projects": [
            {
                "id": "unh", "name": "Urban Nest Heights", "type": "Residential (2/3 BHK apartments)",
                "location": "Gachibowli, Hyderabad", "units": 350, "towers": 3,
                "price_range": "INR 65 lakh - 1.4 crore", "possession": "January 2028",
                "rera": "P01100005421",
            },
            {
                "id": "unr", "name": "Urban Nest Riverside", "type": "Residential (2/3 BHK apartments)",
                "location": "Kokapet, Hyderabad", "units": 410, "towers": 4,
                "price_range": "INR 70 lakh - 1.55 crore", "possession": "August 2028",
                "rera": "P01100005780",
            },
        ],
    },
}

ALL_PROJECTS = [(bk, b, p) for bk, b in BUILDERS.items() for p in b["projects"]]

# ---------------------------------------------------------------------------
# PDF helpers (reportlab)
# ---------------------------------------------------------------------------
styles = getSampleStyleSheet()
h1 = ParagraphStyle("h1", parent=styles["Heading1"], spaceAfter=12)
h2 = ParagraphStyle("h2", parent=styles["Heading2"], spaceAfter=8, spaceBefore=12)
body = ParagraphStyle("body", parent=styles["BodyText"], spaceAfter=8, leading=15)


def write_pdf(path, title, sections):
    """sections: list of (heading, [paragraph, paragraph, ...])"""
    doc = SimpleDocTemplate(path, pagesize=letter,
                             topMargin=54, bottomMargin=54, leftMargin=54, rightMargin=54)
    story = [Paragraph(title, h1), Spacer(1, 6)]
    for heading, paras in sections:
        story.append(Paragraph(heading, h2))
        for p in paras:
            story.append(Paragraph(p, body))
    doc.build(story)


# ---------------------------------------------------------------------------
# DOCX helper
# ---------------------------------------------------------------------------
def write_docx(path, title, sections):
    doc = Document()
    doc.add_heading(title, level=1)
    for heading, paras in sections:
        doc.add_heading(heading, level=2)
        for p in paras:
            doc.add_paragraph(p)
    doc.save(path)


# ---------------------------------------------------------------------------
# HTML / Markdown helpers
# ---------------------------------------------------------------------------
def write_html(path, title, sections):
    parts = [f"<html><head><title>{title}</title></head><body>", f"<h1>{title}</h1>"]
    for heading, paras in sections:
        parts.append(f"<h2>{heading}</h2>")
        for p in paras:
            parts.append(f"<p>{p}</p>")
    parts.append("</body></html>")
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(parts))


def write_md(path, title, sections):
    lines = [f"# {title}", ""]
    for heading, paras in sections:
        lines.append(f"## {heading}")
        lines.append("")
        for p in paras:
            lines.append(p)
            lines.append("")
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


count = 0

# ---------------------------------------------------------------------------
# 1) PROJECT BROCHURES (PDF) - one per project
# ---------------------------------------------------------------------------
for bk, b, p in ALL_PROJECTS:
    sections = [
        ("Project Overview", [
            f"{p['name']} is a {p['type']} project developed by {b['name']}, located in "
            f"{p['location']}. The project spans {p['towers'] if p['towers'] else 'multiple'} "
            f"towers/blocks with a total of {p['units']} units.",
            f"Possession is scheduled for {p['possession']}, subject to regulatory and "
            f"construction timelines. RERA Registration Number: {p['rera']}.",
        ]),
        ("Pricing", [
            f"Indicative pricing for {p['name']} ranges from {p['price_range']}, "
            f"exclusive of GST, registration charges, and other statutory levies.",
            "Final pricing depends on unit type, floor, and facing, and is confirmed at "
            "the time of booking via the official price list.",
        ]),
        ("Specifications", [
            "Structure: RCC framed structure designed to withstand seismic zone requirements.",
            "Flooring: Vitrified tiles in living/dining/bedrooms, anti-skid tiles in "
            "bathrooms and balconies.",
            "Kitchen: Granite platform with stainless steel sink and provision for "
            "modular kitchen fittings.",
            "Electrical: Concealed copper wiring with MCB distribution board, and "
            "provision for AC points in all bedrooms.",
            "Doors and Windows: Engineered wood main door, UPVC/aluminium windows with "
            "mosquito mesh.",
        ]),
        ("Amenities", [
            "Clubhouse with indoor games, gymnasium, and multipurpose hall.",
            "Swimming pool with dedicated kids' pool, landscaped gardens, and jogging track.",
            "24x7 security with CCTV surveillance, intercom facility, and power backup for "
            "common areas and lifts.",
            "Dedicated children's play area, senior citizen sit-out zones, and rainwater "
            "harvesting system.",
        ]),
        ("Connectivity", [
            f"{p['name']} is well connected to major IT corridors, schools, hospitals, and "
            f"shopping centers in and around {p['location']}.",
            "Nearest railway station and airport connectivity details are available on request "
            "from the sales team.",
        ]),
    ]
    write_pdf(os.path.join(PDF_DIR, f"{p['id']}_brochure.pdf"),
              f"{p['name']} — Project Brochure", sections)
    count += 1

# ---------------------------------------------------------------------------
# 2) PAYMENT PLANS (PDF) - one per project (first project per builder to keep count sane)
# ---------------------------------------------------------------------------
for bk, b, p in ALL_PROJECTS:
    sections = [
        ("Construction-Linked Payment Plan (CLP)", [
            "10% of the total consideration value is payable at the time of booking.",
            "15% is payable on execution of the Agreement for Sale, within 30 days of booking.",
            "10% on completion of the foundation/plinth level.",
            "10% on completion of the ground floor slab.",
            "20% spread across completion of each subsequent floor slab (payable in "
            "installments as construction progresses).",
            "10% on completion of internal plastering and flooring works.",
            "10% on completion of external finishing and common area works.",
            "5% at the time of possession/handover, on receipt of Occupancy Certificate.",
        ]),
        ("Down Payment Plan (DPP)", [
            "Buyers opting for the Down Payment Plan are eligible for an additional discount "
            "of 3-5% on the base price, subject to the builder's applicable policy at the "
            "time of booking.",
            "95% of the total consideration is payable within 60 days of booking; the "
            "remaining 5% is payable at possession.",
        ]),
        ("Payment Terms & Conditions", [
            f"All payments for {p['name']} must be made via cheque, demand draft, or bank "
            f"transfer in favor of the designated project escrow account, in compliance "
            f"with RERA (Real Estate Regulatory Authority) escrow account regulations.",
            "Delayed payments beyond the due date will attract interest as specified in the "
            "Agreement for Sale, typically aligned with the State Bank of India's prevailing "
            "MCLR plus an applicable margin.",
            "GST, stamp duty, registration charges, and other statutory levies are payable "
            "additionally as per applicable government rates at the time of each payment.",
        ]),
    ]
    write_pdf(os.path.join(PDF_DIR, f"{p['id']}_payment_plan.pdf"),
              f"{p['name']} — Payment Plan", sections)
    count += 1

# ---------------------------------------------------------------------------
# 3) RERA CERTIFICATE SUMMARY (PDF) - one per project
# ---------------------------------------------------------------------------
for bk, b, p in ALL_PROJECTS:
    sections = [
        ("RERA Registration Details", [
            f"Project Name: {p['name']}",
            f"Promoter: {b['name']}",
            f"RERA Registration Number: {p['rera']}",
            f"Project Location: {p['location']}",
            f"Total Units Registered: {p['units']}",
            f"Proposed Date of Completion: {p['possession']}",
        ]),
        ("Regulatory Compliance", [
            "This project is registered under the Real Estate (Regulation and Development) "
            "Act, 2016 (RERA). 70% of the amounts realized from allottees are maintained in "
            "a designated escrow account to be used only for construction and land costs of "
            "this project, as mandated under Section 4(2)(l)(D) of the Act.",
            "Buyers are advised to verify the current status of this registration on the "
            "respective State RERA authority's official website before making any payment, "
            "since registration validity, extensions, and any regulatory orders are updated "
            "there directly by the authority.",
            "Any structural defect or deficiency in quality reported by an allottee within "
            "five years of possession shall be rectified by the promoter without further "
            "charge, within 30 days of being notified, as per Section 14(3) of the RERA Act.",
        ]),
    ]
    write_pdf(os.path.join(PDF_DIR, f"{p['id']}_rera_summary.pdf"),
              f"{p['name']} — RERA Registration Summary", sections)
    count += 1

# ---------------------------------------------------------------------------
# 4) BUILDER PROFILE (PDF) - one per builder
# ---------------------------------------------------------------------------
for bk, b in BUILDERS.items():
    proj_lines = [f"{p['name']} — {p['type']}, {p['location']} ({p['units']} units, "
                   f"possession {p['possession']})" for p in b["projects"]]
    sections = [
        ("Company Overview", [
            f"{b['name']} was founded in {b['founded']} and is headquartered in {b['hq']}. "
            f"The company has delivered residential and commercial developments across "
            f"multiple cities over more than two decades of operation.",
        ]),
        ("Ongoing Projects", proj_lines),
        ("Track Record", [
            f"{b['name']} has maintained a consistent record of RERA-compliant project "
            f"registrations and has not been subject to any project deregistration by the "
            f"respective State RERA authority as of the latest available filing.",
            "Prospective buyers are encouraged to independently verify project-specific "
            "RERA registration status, litigation history, and completion track record via "
            "the official State RERA portal before booking.",
        ]),
    ]
    write_pdf(os.path.join(PDF_DIR, f"{bk}_builder_profile.pdf"),
              f"{b['name']} — Builder Profile", sections)
    count += 1

print(f"PDFs done: {count}")

# ---------------------------------------------------------------------------
# 5) SALE AGREEMENT / LEGAL TERMS (DOCX) - one per project (first 6 projects)
# ---------------------------------------------------------------------------
docx_count = 0
for bk, b, p in ALL_PROJECTS[:6]:
    sections = [
        ("Definitions", [
            "\"Allottee\" means the person(s) to whom the unit is allotted, sold, or "
            "transferred by the Promoter, and includes their legal heirs and permitted assigns.",
            "\"Promoter\" means " + b["name"] + ", registered under the applicable RERA "
            "authority for this project.",
            "\"Carpet Area\" means the net usable floor area of the unit, excluding the "
            "area covered by external walls, areas under services shafts, exclusive "
            "balcony/verandah, and exclusive open terrace, as defined under the RERA Act.",
        ]),
        ("Terms of Sale", [
            f"This Agreement for Sale governs the sale of Unit(s) in {p['name']} between "
            f"the Promoter and the Allottee, subject to full payment as per the agreed "
            f"payment plan and compliance with all applicable laws.",
            "The Promoter shall not make any additions or alterations in the sanctioned "
            "plans, layout plans, and specifications of the building without the previous "
            "written consent of at least two-thirds of the allottees, other than minor "
            "alterations required due to architectural and structural reasons.",
            "Title to the unit, along with the proportionate undivided share in common "
            "areas, shall be conveyed to the Allottee upon full payment and execution of "
            "the sale/conveyance deed.",
        ]),
        ("Allottee Obligations", [
            "The Allottee shall make timely payments as per the payment plan and shall be "
            "liable to pay interest for any delay, at the rate prescribed under the "
            "applicable RERA rules.",
            "The Allottee shall participate in the formation of the association of "
            "allottees/society/federation, as applicable, and shall pay maintenance charges "
            "as levied from the date of offer of possession.",
        ]),
        ("Dispute Resolution", [
            "Any dispute arising out of this Agreement shall first be referred to mutual "
            "discussion between the parties, failing which either party may approach the "
            "concerned State Real Estate Regulatory Authority or Adjudicating Officer, or "
            "avail remedies under applicable consumer protection law.",
        ]),
    ]
    write_docx(os.path.join(DOCX_DIR, f"{p['id']}_sale_agreement_terms.docx"),
               f"{p['name']} — Agreement for Sale: Key Terms", sections)
    docx_count += 1

# ---------------------------------------------------------------------------
# 6) REGISTRATION PROCESS GUIDE (DOCX) - one per project (first 6)
# ---------------------------------------------------------------------------
for bk, b, p in ALL_PROJECTS[:6]:
    sections = [
        ("Documents Required", [
            "PAN Card and Aadhaar Card (or valid government-issued photo ID) of all buyers.",
            "Passport-size photographs of all buyers and, where applicable, of the seller.",
            "Copy of the executed Agreement for Sale / Sale Deed.",
            "Proof of payment of stamp duty and registration fee (challan/receipt).",
            "No Objection Certificate (NOC) from the financing bank, if the unit is under a "
            "home loan.",
        ]),
        ("Step-by-Step Registration Process", [
            f"1. Compute applicable stamp duty and registration charges based on the "
            f"circle rate/guidance value of the property in {p['location']} and the "
            f"declared sale consideration, whichever is higher.",
            "2. Pay stamp duty online through the respective State's e-stamping portal, or "
            "through an authorized franking center.",
            "3. Book an appointment slot with the local Sub-Registrar office (many states "
            "offer online slot booking).",
            "4. Both buyer and seller (or their authorized representatives with valid Power "
            "of Attorney) must be physically present at the Sub-Registrar office along with "
            "two witnesses.",
            "5. Biometric verification and document execution take place at the "
            "Sub-Registrar office, after which the registered document is issued, typically "
            "within a few working days.",
        ]),
        ("Charges", [
            "Stamp duty and registration charges vary by state and are subject to change "
            "as per the prevailing government notification at the time of registration. "
            "Buyers should confirm current rates with the sales/CRM team or the Sub-Registrar "
            "office directly before budgeting for these costs.",
        ]),
    ]
    write_docx(os.path.join(DOCX_DIR, f"{p['id']}_registration_process.docx"),
               f"{p['name']} — Registration Process Guide", sections)
    docx_count += 1

# ---------------------------------------------------------------------------
# 7) POSSESSION GUIDELINES (DOCX) - one per project (first 6)
# ---------------------------------------------------------------------------
for bk, b, p in ALL_PROJECTS[:6]:
    sections = [
        ("Possession Timeline", [
            f"The proposed date of possession for {p['name']} is {p['possession']}, as "
            f"declared in the RERA registration (Registration No. {p['rera']}). This date "
            f"is subject to a grace period as specified in the Agreement for Sale, "
            f"typically up to 6 months, to account for force majeure events.",
        ]),
        ("Pre-Possession Checklist", [
            "Ensure all payment installments as per the payment plan, including any "
            "applicable interest for delayed payments, have been cleared in full.",
            "Complete registration of the Sale Deed / Conveyance Deed before or at the time "
            "of taking possession, as applicable under the Agreement.",
            "Carry a valid photo ID and the original Agreement for Sale / Allotment Letter "
            "to the possession/handover appointment.",
        ]),
        ("Snag Inspection & Handover", [
            "Buyers are entitled to a joint inspection of the unit with the builder's "
            "representative prior to formal handover, during which any visible defects, "
            "incomplete fittings, or discrepancies from agreed specifications should be "
            "recorded in a joint snag list.",
            "The Promoter is obligated under Section 14(3) of the RERA Act to rectify "
            "structural defects or deficiencies reported by the Allottee within 5 years of "
            "possession, at no additional cost, within 30 days of being notified.",
        ]),
        ("Occupancy Certificate", [
            "Possession will only be offered after the project (or the relevant "
            "tower/block) has received a valid Occupancy Certificate (OC) from the "
            "competent local municipal/planning authority. Buyers may request a copy of "
            "the OC from the builder at the time of possession.",
        ]),
    ]
    write_docx(os.path.join(DOCX_DIR, f"{p['id']}_possession_guidelines.docx"),
               f"{p['name']} — Possession Guidelines", sections)
    docx_count += 1

print(f"DOCX done: {docx_count}")

# ---------------------------------------------------------------------------
# 8) CANCELLATION & REFUND POLICY (DOCX) - one per builder
# ---------------------------------------------------------------------------
for bk, b in BUILDERS.items():
    sections = [
        ("Cancellation by the Allottee", [
            f"An Allottee may request cancellation of their booking with {b['name']} by "
            f"submitting a written request to the sales/CRM team, along with the original "
            f"booking receipt and payment proofs.",
            "If cancellation is requested within 15 days of booking (the free-look period), "
            "the entire amount paid shall be refunded after deducting only the actual bank "
            "transaction charges, if any.",
            "If cancellation is requested after the free-look period but before execution "
            "of the Agreement for Sale, an amount not exceeding the earnest money "
            "(typically 10% of the total consideration) shall be forfeited, and the balance "
            "refunded within 45 days of the cancellation request.",
            "If cancellation is requested after execution of the Agreement for Sale, the "
            "forfeiture amount and refund timeline shall be governed by the specific clauses "
            "of the executed Agreement for Sale, consistent with applicable RERA rules "
            "capping forfeiture amounts.",
        ]),
        ("Cancellation by the Promoter", [
            f"{b['name']} may cancel an allotment only for valid reasons specified in the "
            f"Agreement for Sale, such as non-payment of dues after due notice, and shall "
            f"refund the amount paid by the Allottee after lawful deductions, with interest "
            f"as applicable under the RERA Act, within 45 days of such cancellation.",
        ]),
        ("Refund Processing", [
            "All approved refunds are processed via bank transfer (NEFT/RTGS) to the "
            "original payment account of the Allottee within 30-45 working days of refund "
            "approval, subject to completion of any pending documentation.",
            "Refunds are not applicable for amounts paid towards statutory charges (GST, "
            "stamp duty already deposited with government authorities, TDS) once remitted "
            "to the relevant government authority.",
        ]),
    ]
    write_docx(os.path.join(DOCX_DIR, f"{bk}_cancellation_refund_policy.docx"),
               f"{b['name']} — Cancellation & Refund Policy", sections)
    docx_count += 1

print(f"DOCX total: {docx_count}")

# ---------------------------------------------------------------------------
# 9) BUILDER WEBSITE PAGES (HTML): Home, About, Privacy, Terms, FAQ — one set per builder
# ---------------------------------------------------------------------------
html_count = 0
for bk, b in BUILDERS.items():
    proj_lines = [f"{p['name']} ({p['type']}) in {p['location']} — {p['price_range']}"
                  for p in b["projects"]]

    write_html(os.path.join(HTML_DIR, f"{bk}_home.html"), f"{b['name']} — Home", [
        ("Welcome", [
            f"{b['name']} is a real estate developer headquartered in {b['hq']}, "
            f"delivering residential and commercial projects since {b['founded']}.",
        ]),
        ("Our Projects", proj_lines),
        ("Why Choose Us", [
            "RERA-registered projects across all active developments.",
            "Transparent construction-linked payment plans.",
            "Dedicated customer support for the entire ownership lifecycle, from booking "
            "to post-possession service requests.",
        ]),
    ])
    html_count += 1

    write_html(os.path.join(HTML_DIR, f"{bk}_about.html"), f"About {b['name']}", [
        ("Company Profile", [
            f"Founded in {b['founded']} and headquartered in {b['hq']}, {b['name']} has "
            f"grown into a recognized name in the region's real estate sector, with a "
            f"portfolio spanning residential apartments, villas, and commercial spaces.",
        ]),
        ("Leadership & Governance", [
            f"{b['name']} operates under a professional management structure with "
            f"dedicated legal, engineering, and customer relationship teams overseeing "
            f"each project from land acquisition through possession.",
        ]),
        ("Certifications", [
            "All active projects are registered with the respective State Real Estate "
            "Regulatory Authority (RERA) prior to launch and marketing, in compliance with "
            "the RERA Act, 2016.",
        ]),
    ])
    html_count += 1

    write_html(os.path.join(HTML_DIR, f"{bk}_privacy_policy.html"),
               f"{b['name']} — Privacy Policy", [
        ("Information We Collect", [
            "We collect personal information you provide directly, such as your name, "
            "phone number, email address, and address, when you enquire about a project, "
            "book a site visit, or register interest through our website or sales team.",
            "We may also collect information automatically through cookies and analytics "
            "tools when you browse our website, including IP address, browser type, and "
            "pages visited.",
        ]),
        ("How We Use Your Information", [
            "To respond to enquiries, schedule site visits, and share project brochures, "
            "pricing, and payment plan details.",
            "To send updates regarding construction progress, possession timelines, and "
            "regulatory filings related to a project you have booked.",
            "We do not sell your personal information to third parties. We may share "
            "necessary details with our empaneled banking/home loan partners only with your "
            "explicit consent.",
        ]),
        ("Your Rights", [
            "You may request access to, correction of, or deletion of your personal data "
            "held by us by contacting our customer support team, subject to any statutory "
            "retention obligations (e.g. records required to be maintained under RERA or "
            "tax law).",
        ]),
    ])
    html_count += 1

    write_html(os.path.join(HTML_DIR, f"{bk}_terms_conditions.html"),
               f"{b['name']} — Terms & Conditions", [
        ("Website Use", [
            f"These Terms & Conditions govern your use of the {b['name']} website. By "
            f"accessing this website, you agree to be bound by these terms.",
            "All project information, images, floor plans, and pricing on this website are "
            "indicative and subject to change without prior notice. The Agreement for Sale "
            "executed at the time of booking is the final and binding document.",
        ]),
        ("Booking Terms", [
            "A booking is confirmed only upon receipt of the booking amount and issuance of "
            "a formal Allotment Letter/Booking Confirmation by the Promoter.",
            "All bookings are subject to availability of the unit at the time of payment "
            "realization, regardless of any prior verbal confirmation.",
        ]),
        ("Limitation of Liability", [
            f"{b['name']} shall not be liable for any indirect or consequential loss "
            f"arising from reliance on indicative information published on this website, "
            f"to the extent permitted by applicable law.",
        ]),
    ])
    html_count += 1

    write_html(os.path.join(HTML_DIR, f"{bk}_faq.html"),
               f"{b['name']} — Frequently Asked Questions", [
        ("Booking & Payments", [
            "Q: What is the minimum amount required to book a unit? A: A booking amount, "
            "typically 10% of the total consideration value, is required to confirm a "
            "booking, followed by the construction-linked or down payment plan as chosen.",
            "Q: Can I switch from a construction-linked plan to a down payment plan after "
            "booking? A: Plan changes may be permitted within 30 days of booking, subject "
            "to written request and approval by the sales team; applicable discounts may "
            "differ.",
        ]),
        ("Home Loans", [
            "Q: Is home loan assistance available? A: Yes, our projects are approved by "
            "several leading banks and NBFCs for home loan disbursement; our CRM team can "
            "share the list of approved lenders on request.",
        ]),
        ("Possession & Documentation", [
            "Q: When will I receive possession? A: Possession is offered as per the date "
            "declared in the project's RERA registration, subject to the grace period "
            "specified in the Agreement for Sale.",
            "Q: What documents will I receive at possession? A: You will receive the "
            "possession letter, copy of the Occupancy Certificate, and the completion "
            "snag/handover checklist signed jointly.",
        ]),
    ])
    html_count += 1

print(f"HTML builder pages done: {html_count}")

# ---------------------------------------------------------------------------
# 10) LISTING PORTAL PAGES (HTML) - third-party portal listings + portal policies
# ---------------------------------------------------------------------------
for bk, b, p in ALL_PROJECTS:
    write_html(os.path.join(HTML_DIR, f"propertybazaar_listing_{p['id']}.html"),
               f"{p['name']} — Listing on PropertyBazaar", [
        ("Listing Summary", [
            f"{p['name']} by {b['name']} is listed on PropertyBazaar under "
            f"{p['type']}, located at {p['location']}.",
            f"Price range as listed: {p['price_range']}. Total units: {p['units']}. "
            f"Possession: {p['possession']}.",
        ]),
        ("User Reviews Summary", [
            "Aggregated user ratings on PropertyBazaar for this listing reflect feedback "
            "primarily on construction quality, connectivity, and adherence to the "
            "declared possession timeline; individual reviews are visible on the live "
            "listing page and are not reproduced in this document.",
        ]),
        ("Disclaimer", [
            "Listing details are provided by the builder/broker and periodically verified "
            "by PropertyBazaar; buyers should independently confirm RERA registration, "
            "pricing, and specifications directly with the builder before making any "
            "payment.",
        ]),
    ])
    html_count += 1

write_html(os.path.join(HTML_DIR, "propertybazaar_terms_of_use.html"),
           "PropertyBazaar — Terms of Use", [
    ("Use of Platform", [
        "PropertyBazaar is a listing platform that aggregates project and resale property "
        "information from builders, brokers, and individual owners. PropertyBazaar is not "
        "a party to any transaction between buyers and sellers/builders listed on the "
        "platform.",
    ]),
    ("Accuracy of Listings", [
        "While PropertyBazaar makes reasonable efforts to verify listing information, "
        "pricing, availability, and RERA registration status may change without notice. "
        "Users are advised to independently verify all details with the listing party "
        "before making any payment or commitment.",
    ]),
])
html_count += 1

write_html(os.path.join(HTML_DIR, "propertybazaar_privacy_policy.html"),
           "PropertyBazaar — Privacy Policy", [
    ("Data Collection", [
        "PropertyBazaar collects contact details submitted by users when enquiring about a "
        "listing, and may share these details with the relevant builder or broker to "
        "facilitate a response to the enquiry.",
    ]),
    ("Data Sharing", [
        "PropertyBazaar does not sell user data to unrelated third parties. Contact details "
        "shared with a listed builder/broker are governed additionally by that builder's "
        "own privacy policy once received.",
    ]),
])
html_count += 1

print(f"HTML total: {html_count}")

# ---------------------------------------------------------------------------
# 11) HOME LOAN INFORMATION (Markdown)
# ---------------------------------------------------------------------------
md_count = 0
write_md(os.path.join(MD_DIR, "home_loan_general_guide.md"),
         "Home Loan Information Guide", [
    ("Eligibility", [
        "- Salaried applicants: minimum age 21 years, stable income with at least 2 years "
        "of continuous employment, typically required.",
        "- Self-employed applicants: minimum 3 years of business continuity and audited "
        "financial statements for the last 2-3 years typically required.",
        "- Loan eligibility is generally computed based on a debt-to-income ratio, with "
        "EMI (including the proposed home loan) usually capped around 40-50% of net "
        "monthly income, though this varies by lender.",
    ]),
    ("Documents Typically Required", [
        "- KYC documents: PAN card, Aadhaar card, passport-size photographs.",
        "- Income proof: salary slips (last 3-6 months) and Form 16 / ITR for salaried "
        "applicants; ITR and audited financials for self-employed applicants.",
        "- Property documents: copy of the Agreement for Sale, RERA registration "
        "certificate, approved building plan, and title documents (to be provided by the "
        "builder's CRM team on request).",
        "- Bank statements for the last 6 months.",
    ]),
    ("Loan-to-Value (LTV) Ratio", [
        "- Properties valued up to INR 30 lakh: up to 90% LTV typically permissible.",
        "- Properties valued INR 30 lakh - 75 lakh: up to 80% LTV typically permissible.",
        "- Properties valued above INR 75 lakh: up to 75% LTV typically permissible.",
        "- These are general regulatory ceilings; actual sanctioned LTV depends on the "
        "specific lender's credit policy and the applicant's credit profile.",
    ]),
    ("Disbursement Against Construction-Linked Plans", [
        "- For under-construction projects, loan disbursement is typically released in "
        "tranches aligned with the construction-linked payment plan milestones, upon "
        "submission of the builder's demand letter for each stage.",
        "- Pre-EMI interest (interest-only payments on the disbursed amount) is usually "
        "charged until full disbursement, after which regular EMI (principal + interest) "
        "begins.",
    ]),
])
md_count += 1

for bk, b, p in ALL_PROJECTS[:4]:
    write_md(os.path.join(MD_DIR, f"{p['id']}_home_loan_partners.md"),
             f"{p['name']} — Home Loan Partner Information", [
        ("Approved Lenders", [
            f"{p['name']} is pre-approved for home loan disbursement by several major "
            f"nationalized and private banks as well as leading housing finance "
            f"companies. The current list of approved lenders is maintained by the CRM "
            f"team and shared at the time of booking, since bank empanelment can change "
            f"during the construction period.",
        ]),
        ("Process", [
            f"1. Buyer selects a preferred lender from the approved list (or applies to an "
            f"external lender, subject to that lender's own project appraisal).",
            "2. Buyer submits KYC, income, and property documents (Agreement for Sale, "
            "RERA certificate, approved plan) to the lender.",
            "3. Lender conducts technical and legal due diligence on the project before "
            "sanctioning the loan.",
            f"4. Disbursement is released in tranches matching {p['name']}'s "
            f"construction-linked payment plan milestones.",
        ]),
    ])
    md_count += 1

print(f"Markdown home loan docs done: {md_count}")

# ---------------------------------------------------------------------------
# 12) CUSTOMER SUPPORT DOCUMENTATION (Markdown) - one per builder
# ---------------------------------------------------------------------------
for bk, b in BUILDERS.items():
    write_md(os.path.join(MD_DIR, f"{bk}_customer_support.md"),
             f"{b['name']} — Customer Support Documentation", [
        ("Support Channels", [
            "- Phone support: available Monday-Saturday, 9:30 AM - 6:30 PM, for booking, "
            "payment, and general project queries.",
            "- Email support: for documentation requests, payment receipts, and "
            "post-possession service requests, with a standard response time of 2 "
            "business days.",
            "- Site visit desk: available at each project site during sales office hours "
            "for walk-in queries and scheduled visits.",
        ]),
        ("Common Request Types", [
            "- Payment receipt / statement of account requests.",
            "- Demand letter clarification for construction-linked payment installments.",
            "- Post-possession maintenance and snag/defect reporting.",
            "- Documentation requests: Allotment Letter copy, Agreement for Sale copy, "
            "RERA registration certificate copy.",
        ]),
        ("Escalation Process", [
            f"If a query is not resolved by the initial support channel within the stated "
            f"turnaround time, it may be escalated to the project's CRM manager, and "
            f"subsequently to {b['name']}'s central customer relations team. Unresolved "
            f"grievances relating to RERA-covered matters may also be filed with the "
            f"respective State Real Estate Regulatory Authority.",
        ]),
    ])
    md_count += 1

# ---------------------------------------------------------------------------
# 13) AMENITIES GUIDES (Markdown) - one per project (first 3)
# ---------------------------------------------------------------------------
for bk, b, p in ALL_PROJECTS[:6]:
    write_md(os.path.join(MD_DIR, f"{p['id']}_amenities_guide.md"),
             f"{p['name']} — Amenities Guide", [
        ("Recreational Amenities", [
            "- Clubhouse with indoor games room, gymnasium, and multipurpose banquet hall.",
            "- Swimming pool with a separate children's pool and deck seating area.",
            "- Landscaped gardens, jogging/walking track, and outdoor seating pockets.",
            "- Dedicated children's play area with age-appropriate play equipment.",
        ]),
        ("Safety & Convenience", [
            "- 24x7 security personnel with CCTV coverage at entry/exit points and common "
            "areas.",
            "- Video door phone / intercom connectivity to the main gate.",
            "- Power backup for common areas, lifts, and (in most unit types) a defined "
            "backup load per apartment.",
            "- Rainwater harvesting and segregated waste management (wet/dry) as per local "
            "municipal norms.",
        ]),
        ("Sustainability Features", [
            "- Solar-powered lighting for common outdoor areas where applicable.",
            "- Water treatment plant / STP (Sewage Treatment Plant) for treated water reuse "
            "in landscaping.",
        ]),
    ])
    md_count += 1

# ---------------------------------------------------------------------------
# 14) LOCATION GUIDES (Markdown) - one per project (first 3)
# ---------------------------------------------------------------------------
for bk, b, p in ALL_PROJECTS[:6]:
    write_md(os.path.join(MD_DIR, f"{p['id']}_location_guide.md"),
             f"{p['name']} — Location & Connectivity Guide", [
        ("Overview", [
            f"{p['name']} is located in {p['location']}, an area known for its proximity "
            f"to established residential neighborhoods, commercial hubs, and social "
            f"infrastructure.",
        ]),
        ("Connectivity", [
            "- Well connected via arterial roads to the nearest IT/business corridor, "
            "typically within a 15-25 minute drive depending on traffic conditions.",
            "- Public transport access via nearby bus routes; metro/rail connectivity "
            "details, where applicable, are updated periodically as transit infrastructure "
            "expands in the region.",
        ]),
        ("Social Infrastructure", [
            "- Reputed schools and colleges within a 3-5 km radius.",
            "- Multi-specialty hospitals and diagnostic centers in the vicinity.",
            "- Shopping malls, supermarkets, and everyday retail within easy reach of the "
            "project.",
        ]),
    ])
    md_count += 1

# ---------------------------------------------------------------------------
# 15) FLOOR PLAN DESCRIPTIONS (Markdown) - one per project (first 3)
# ---------------------------------------------------------------------------
for bk, b, p in ALL_PROJECTS[:6]:
    write_md(os.path.join(MD_DIR, f"{p['id']}_floor_plans.md"),
             f"{p['name']} — Floor Plan Descriptions", [
        ("Unit Types", [
            f"{p['name']} offers {p['type']}. Exact carpet area and configuration vary by "
            f"tower and floor; the figures below are representative of the typical unit "
            f"mix offered at launch.",
        ]),
        ("Typical Configurations", [
            "- 2 BHK: approx. 950-1150 sq. ft. carpet area, 2 bedrooms, 2 bathrooms, "
            "living/dining area, kitchen, and a balcony off the living room.",
            "- 3 BHK: approx. 1350-1650 sq. ft. carpet area, 3 bedrooms (including one "
            "master with attached bathroom), 2-3 bathrooms, living/dining, kitchen with "
            "utility area, and balconies off the living room and master bedroom.",
        ]),
        ("Notes", [
            "Actual carpet area, as defined under RERA, is stated precisely in the "
            "Agreement for Sale for the specific unit booked and may vary from the "
            "brochure's indicative figures, which are subject to minor changes during "
            "detailed structural design.",
        ]),
    ])
    md_count += 1

# ---------------------------------------------------------------------------
# 16) RERA GENERAL INFORMATION (Markdown)
# ---------------------------------------------------------------------------
write_md(os.path.join(MD_DIR, "rera_general_information.md"),
         "Understanding RERA — General Information for Homebuyers", [
    ("What is RERA", [
        "The Real Estate (Regulation and Development) Act, 2016 (RERA) is a central "
        "legislation aimed at protecting homebuyers and boosting investment in the real "
        "estate sector by mandating transparency and accountability from developers.",
    ]),
    ("Key Buyer Protections", [
        "- Mandatory registration of projects above a specified size threshold with the "
        "State RERA authority before advertising, marketing, booking, or selling any unit.",
        "- 70% of funds collected from buyers must be maintained in a separate escrow "
        "account, used only for construction and land costs of that specific project.",
        "- Developers must disclose the carpet area (not super built-up area) as the "
        "standard basis for pricing and cannot change sanctioned plans without two-thirds "
        "consent of allottees.",
        "- Structural defects reported within 5 years of possession must be rectified by "
        "the promoter free of charge within 30 days of being notified.",
        "- Delayed possession entitles the buyer to either a refund with interest or "
        "continued possession of the unit with monthly interest compensation, at the "
        "buyer's option.",
    ]),
    ("How to Verify a Project", [
        "Every buyer can verify a project's registration number, promoter details, "
        "layout approval, and complaint history directly on the respective State RERA "
        "authority's official website before making any booking payment.",
    ]),
])
md_count += 1

print(f"Markdown total: {md_count}")

total_docs = count + docx_count + html_count + md_count
print(f"\nGRAND TOTAL DOCUMENTS: {total_docs}")
print(f"  PDF: {count}")
print(f"  DOCX: {docx_count}")
print(f"  HTML: {html_count}")
print(f"  Markdown: {md_count}")
